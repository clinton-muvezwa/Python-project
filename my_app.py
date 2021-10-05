from docx import Document
from docx.shared import Inches
import pyttsx3

# pyttsx3.speak('Hello')
def speak(text):
    pyttsx3.speak(text)

document = Document()

# profile picture
document.add_picture(
    'clint.jpg',
    width=Inches(1.2)
)

# name and contact details
name = input('What is your name?: ')
speak("Hello" + name + " " + "how are you?")
speak('Please enter your details below!')
phone_number = input('What is your phone number?: ')
email = input('What is your email?: ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# about me
document.add_heading('About me')
p = document.add_paragraph()
speak("Please tell me about yourself, press enter to finish")
about_me = input('Tell me about yourself?: ')
p.add_run(about_me + '\n')
# alternative
# document.add_paragraph(
# input('Tell me about yourself?: ')
# )
while True:
    about_me = input(
        'More about yourself? enter "Return" when done: ')
    if about_me == '':
        break
    else:
        p.add_run(about_me + '\n')

# work experience
document.add_heading('Work experience')
print('Your work experience ' + '\n')
p = document.add_paragraph()

company = input('Enter company: ')
from_date = input('From date: ')
to_date = input('To date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_details = input(
    'Describe your experience at ' + company + ': ')

p.add_run(experience_details)

while True:
    add_more_experiences = input(
        'Do you want to add more experiences? Yes or No: ')
    if add_more_experiences.lower() == 'yes':
        p = document.add_paragraph()
        company = input('Enter company ')
        from_date = input('From date: ')
        to_date = input('To date: ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your experience at ' + company + ': ')

        p.add_run(experience_details)
    else:
        break

# Skills
document.add_heading('Skills')
print('Your skills ' + '\n')
p = document.add_paragraph()
p.style = 'List Bullet'

skill = input('Enter skill: ')
years_of_experience = input('Years: ')

p.add_run(skill + ' ').bold = True
p.add_run(years_of_experience + ' yrs' + '\n').italic = True

skill_details = input(
    'Summarise your ' + skill + ' usage: ')

p.add_run(skill_details)

while True:
    add_more_skills = input(
        'Do you want to add more skills? Yes or No: ')
    if add_more_skills.lower() == 'yes':
        p = document.add_paragraph()
        p.style = 'List Bullet'

        skill = input('Enter skill: ')
        years_of_experience = input('Years: ')

        p.add_run(skill + ' ').bold = True
        p.add_run(years_of_experience + ' yrs' + '\n').italic = True

        skill_details = input(
            'Summarise your ' + skill + ' usage: ')

        p.add_run(skill_details)

    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Python"

document.save('clint_cv.docx')

